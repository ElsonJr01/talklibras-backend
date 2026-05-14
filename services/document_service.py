import cloudinary
import cloudinary.uploader
import cloudinary.api
import os
import re
import uuid
import httpx
from datetime import datetime
import PyPDF2
import docx
import io
from dotenv import load_dotenv

load_dotenv()

# Cloudinary config
cloudinary.config(
    cloud_name=os.getenv("CLOUDINARY_CLOUD_NAME"),
    api_key=os.getenv("CLOUDINARY_API_KEY"),
    api_secret=os.getenv("CLOUDINARY_API_SECRET")
)

PLAN_LIMITS = {
    "free": 3,
    "pro": 30,
    "pro_max": 999999
}

def generate_serial_code():
    """Gera código serial único para documentos PRO"""
    raw = str(uuid.uuid4()).upper().replace("-", "")
    return f"TLK-{raw[:4]}-{raw[4:8]}-{raw[8:12]}"

async def upload_to_cloudinary(file_bytes: bytes, filename: str, folder: str = "TalkLib") -> dict:
    """Faz upload do arquivo para Cloudinary"""
    result = cloudinary.uploader.upload(
        file_bytes,
        folder=folder,
        resource_type="raw",
        public_id=f"{folder}/{uuid.uuid4()}_{filename}",
        use_filename=True,
        unique_filename=True
    )
    return {
        "url": result["secure_url"],
        "public_id": result["public_id"],
        "bytes": result.get("bytes", 0)
    }

async def extract_text_from_pdf(file_bytes: bytes) -> tuple[str, int]:
    """Extrai texto de arquivo PDF"""
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text.strip(), len(reader.pages)

async def extract_text_from_docx(file_bytes: bytes) -> tuple[str, int]:
    """Extrai texto de arquivo DOCX"""
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    return text.strip(), len(doc.paragraphs)

async def extract_text_from_txt(file_bytes: bytes) -> tuple[str, int]:
    """Extrai texto de arquivo TXT"""
    text = file_bytes.decode("utf-8", errors="ignore")
    return text.strip(), 1

async def extract_text(file_bytes: bytes, filename: str) -> tuple[str, int, str]:
    """Extrai texto baseado no tipo do arquivo"""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    
    if ext == "pdf":
        text, pages = await extract_text_from_pdf(file_bytes)
        return text, pages, "pdf"
    elif ext in ["docx", "doc"]:
        text, pages = await extract_text_from_docx(file_bytes)
        return text, pages, "docx"
    elif ext == "txt":
        text, pages = await extract_text_from_txt(file_bytes)
        return text, pages, "txt"
    else:
        # Tenta como texto puro
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
            return text.strip(), 1, ext
        except:
            raise ValueError(f"Formato de arquivo não suportado: {ext}")

# ─── TTS Services ────────────────────────────────────────────────────────────

async def tts_google_free(text: str, lang: str = "pt-BR") -> str:
    """
    Usa a API do Google Text-to-Speech (via endpoint público do gTTS/goo.gl)
    Retorna URL do arquivo de áudio
    """
    # Google TTS endpoint público (sem chave - limite por IP)
    base_url = "https://translate.google.com/translate_tts"
    text_chunks = chunk_text(text, 200)
    
    # Retorna URL para o primeiro chunk (frontend faz stream dos demais)
    params = {
        "ie": "UTF-8",
        "q": text_chunks[0] if text_chunks else text[:200],
        "tl": lang.split("-")[0],
        "client": "tw-ob",
        "ttsspeed": "1"
    }
    query = "&".join(f"{k}={v}" for k, v in params.items())
    return f"{base_url}?{query}", text_chunks

async def tts_voicerss_free(text: str, lang: str = "pt-br") -> dict:
    """
    Usa VoiceRSS - gratuito até 350 requisições/dia
    Retorna URL do áudio
    """
    api_key = os.getenv("VOICERSS_API_KEY", "")
    if not api_key:
        return {"error": "VoiceRSS API key não configurada", "url": None}
    
    url = "https://api.voicerss.org/"
    params = {
        "key": api_key,
        "hl": lang,
        "src": text[:500],
        "c": "MP3",
        "f": "44khz_16bit_mono",
        "r": "0",
        "b64": "true"
    }
    async with httpx.AsyncClient() as client:
        resp = await client.get(url, params=params, timeout=15)
        if resp.status_code == 200:
            return {"audio_b64": resp.text, "format": "mp3"}
    return {"error": "Falha no VoiceRSS", "url": None}

async def tts_responsivevoice_url(text: str, voice: str = "Brazilian Portuguese Female") -> str:
    """
    ResponsiveVoice - TTS gratuito para uso público (sem chave para texto simples)
    """
    import urllib.parse
    encoded = urllib.parse.quote(text[:500])
    return f"https://code.responsivevoice.org/getaudio?tl=pt&sv=&vn={urllib.parse.quote(voice)}&pitch=0.5&rate=0.5&vol=1&t=1&text={encoded}"

def chunk_text(text: str, max_chars: int = 200) -> list:
    """Divide texto em chunks para TTS"""
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks = []
    current = ""
    for sentence in sentences:
        if len(current) + len(sentence) <= max_chars:
            current += " " + sentence
        else:
            if current:
                chunks.append(current.strip())
            current = sentence
    if current:
        chunks.append(current.strip())
    return chunks if chunks else [text[:max_chars]]

# ─── Libras Services ──────────────────────────────────────────────────────────

async def get_vlibras_url(text: str) -> dict:
    """
    VLibras - Tradução para Libras (API oficial do governo brasileiro)
    https://vlibras.gov.br/
    """
    # VLibras Widget endpoint
    return {
        "service": "VLibras",
        "embed_url": "https://vlibras.gov.br/app",
        "widget_script": "https://vlibras.gov.br/app/vlibras-plugin.js",
        "text": text,
        "instructions": "Use o widget VLibras para tradução em tempo real",
        "direct_url": f"https://vlibras.gov.br/app?vlibras=&text={text[:200]}"
    }

async def get_hand_talk_info() -> dict:
    """
    Hand Talk - Informações sobre API de Libras
    """
    return {
        "service": "Hand Talk",
        "website": "https://www.handtalk.me",
        "free_tier": "Trial disponível",
        "note": "Requer cadastro para uso"
    }

async def get_libras_signs_for_word(word: str) -> dict:
    """
    Dicionário aberto de sinais de Libras via Spread the Sign
    """
    clean_word = word.lower().strip()
    return {
        "word": clean_word,
        "spread_the_sign_url": f"https://www.spreadthesign.com/pt.br/search/?q={clean_word}",
        "dicionario_libras_url": f"https://dicionariolibras.com.br/palavra/{clean_word}",
        "note": "Links para dicionários abertos de Libras"
    }
